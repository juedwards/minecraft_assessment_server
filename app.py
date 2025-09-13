#!/usr/bin/env python3
"""
Minecraft Education - 3D Tracker with AI Assessment
"""

import asyncio
import json
import logging
import websockets
from uuid import uuid4
from datetime import datetime, timezone
import socket
from http.server import HTTPServer, BaseHTTPRequestHandler
import threading
import time
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
logger = logging.getLogger(__name__)

# Try to import Azure OpenAI with fallback
try:
    from openai import AzureOpenAI
    client = AzureOpenAI(
        api_version=os.getenv('AZURE_OPENAI_API_VERSION', '2025-01-01-preview'),
        azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
        api_key=os.getenv('AZURE_OPENAI_API_KEY'),
    )
    logger.info(f"✅ Azure OpenAI client initialized with endpoint: {os.getenv('AZURE_OPENAI_ENDPOINT')}")
    logger.info(f"✅ Using deployment: {os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME', 'gpt-4.1')}")
except ImportError:
    # Fallback for older openai versions
    import openai
    openai.api_type = "azure"
    openai.api_key = os.getenv('AZURE_OPENAI_API_KEY')
    openai.api_base = os.getenv('AZURE_OPENAI_ENDPOINT')
    openai.api_version = os.getenv('AZURE_OPENAI_API_VERSION', '2025-01-01-preview')
    client = None  # We'll handle this in the analyze function
    logger.info("⚠️  Using fallback OpenAI client (older version)")

# Global storage
player_positions = {}
web_clients = set()
block_events = []
active_players = set()  # Track active players
latest_assessment_results = {}  # Store latest AI assessment results

# Event recording
session_events = []
session_start_time = None
session_id = None
session_file = None
event_buffer = []  # Buffer for batch writing
last_save_time = time.time()

# Data directory
DATA_DIR = "data"

def ensure_data_directory():
    """Ensure the data directory exists"""
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
        logger.info(f"📁 Created data directory: {DATA_DIR}")

def get_external_ip():
    """Get the external IP address of this device"""
    try:
        # Try to connect to an external server to get our IP
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as s:
            s.connect(("8.8.8.8", 80))
            return s.getsockname()[0]
    except:
        # Fallback to hostname
        try:
            hostname = socket.gethostname()
            return socket.gethostbyname(hostname)
        except:
            return "localhost"

def start_session():
    """Start a new recording session"""
    global session_start_time, session_id, session_events, session_file
    session_start_time = datetime.now(timezone.utc)
    session_id = f"minecraft_session_{session_start_time.strftime('%Y%m%d_%H%M%S')}"
    session_events = []
    
    # Record session start event
    initial_event = {
        "timestamp": session_start_time.isoformat(),
        "event_type": "session_start",
        "data": {
            "session_id": session_id,
            "server_version": "1.0",
            "user": "juedwards"
        }
    }
    session_events.append(initial_event)
    
    # Create initial JSON file in data directory
    session_file = os.path.join(DATA_DIR, f"{session_id}.json")
    save_session_realtime()
    
    logger.info(f"📝 Started new session: {session_id}")
    logger.info(f"💾 Real-time saving to: {session_file}")

def save_session_realtime():
    """Save the current session to JSON file (used during active session)"""
    global session_events, session_start_time, session_id, session_file
    
    if not session_file:
        return
    
    current_time = datetime.now(timezone.utc)
    
    # Create session data structure
    session_data = {
        "session_info": {
            "id": session_id,
            "start_time": session_start_time.isoformat(),
            "last_update": current_time.isoformat(),
            "duration_seconds": (current_time - session_start_time).total_seconds(),
            "total_events": len(session_events),
            "user": "juedwards",
            "status": "active" if active_players else "ended"
        },
        "events": session_events
    }
    
    # Write to temporary file first, then rename (atomic operation)
    temp_file = f"{session_file}.tmp"
    with open(temp_file, 'w', encoding='utf-8') as f:
        json.dump(session_data, f, indent=2)
    
    # Rename temp file to actual file (atomic on most systems)
    if os.path.exists(session_file):
        os.remove(session_file)
    os.rename(temp_file, session_file)

def end_session():
    """Mark session as ended and do final save"""
    global session_events
    
    if not session_id:
        return
    
    session_end_time = datetime.now(timezone.utc)
    
    # Add session end event
    session_events.append({
        "timestamp": session_end_time.isoformat(),
        "event_type": "session_end",
        "data": {
            "session_id": session_id,
            "duration_seconds": (session_end_time - session_start_time).total_seconds(),
            "total_events": len(session_events)
        }
    })
    
    # Final save with ended status
    save_session_realtime()
    
    logger.info(f"💾 Session ended and saved: {session_file} ({len(session_events)} events)")

def record_event(event_type, data):
    """Record an event to the session with real-time saving"""
    global session_events, event_buffer, last_save_time
    
    if session_id:  # Only record if session is active
        event = {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "event_type": event_type,
            "data": data
        }
        session_events.append(event)
        event_buffer.append(event)
        
        # Save every 5 seconds or every 50 events
        current_time = time.time()
        if (current_time - last_save_time > 5) or (len(event_buffer) >= 50):
            save_session_realtime()
            event_buffer = []
            last_save_time = current_time
            logger.debug(f"💾 Auto-saved session with {len(session_events)} events")

def create_assessment_document(analyses):
    """Create a Word document from assessment analyses"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Minecraft Player Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Session: {session_id or "N/A"}')
    doc.add_paragraph()
    
    # Add each player's assessment
    for player_name, analysis in analyses.items():
        # Player heading
        player_heading = doc.add_heading(f'Player: {player_name}', 1)
        
        # Parse and format the analysis text
        # Split by common section markers
        sections = analysis.split('\n\n')
        
        for section in sections:
            if section.strip():
                # Check if it's a heading (usually starts with caps or has colons)
                lines = section.strip().split('\n')
                first_line = lines[0].strip()
                
                # Common heading patterns
                if (first_line.isupper() or 
                    first_line.endswith(':') or 
                    first_line.startswith('##') or
                    any(keyword in first_line.upper() for keyword in ['ASSESSMENT', 'CRITERIA', 'FEEDBACK', 'SUGGESTIONS', 'OVERALL'])):
                    
                    # Add as heading
                    heading_text = first_line.replace('##', '').replace(':', '').strip()
                    doc.add_heading(heading_text, 2)
                    
                    # Add remaining lines as content
                    if len(lines) > 1:
                        content = '\n'.join(lines[1:])
                        p = doc.add_paragraph(content)
                        p.style.font.size = Pt(11)
                else:
                    # Add as regular paragraph
                    p = doc.add_paragraph(section)
                    p.style.font.size = Pt(11)
        
        # Add page break between players (except for the last one)
        if player_name != list(analyses.keys())[-1]:
            doc.add_page_break()
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Generated by Minecraft 3D Tracker with AI Assessment"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

async def analyze_player_data():
    """Analyze current player data against rubric using AI"""
    global latest_assessment_results
    try:
        # Read rubric
        rubric_path = os.path.join(os.path.dirname(__file__), 'rubric.md')
        if not os.path.exists(rubric_path):
            return {"error": "Rubric file not found"}
        
        with open(rubric_path, 'r') as f:
            rubric_content = f.read()
        
        # Prepare player data for analysis
        player_analysis_data = {}
        
        # Extract relevant events for each player
        for event in session_events:
            if event['event_type'] in ['player_position', 'block_placed', 'block_broken', 'player_join', 'player_leave']:
                player_id = event['data'].get('player_id')
                player_name = event['data'].get('player_name', player_id)
                
                if player_name not in player_analysis_data:
                    player_analysis_data[player_name] = {
                        'positions': [],
                        'blocks_placed': [],
                        'blocks_broken': [],
                        'join_time': None,
                        'leave_time': None,
                        'total_distance': 0
                    }
                
                if event['event_type'] == 'player_position':
                    player_analysis_data[player_name]['positions'].append(event['data']['position'])
                elif event['event_type'] == 'block_placed':
                    player_analysis_data[player_name]['blocks_placed'].append({
                        'type': event['data']['block_type'],
                        'position': event['data']['estimated_block_position'],
                        'time': event['timestamp']
                    })
                elif event['event_type'] == 'block_broken':
                    player_analysis_data[player_name]['blocks_broken'].append({
                        'type': event['data']['block_type'],
                        'position': event['data']['estimated_block_position'],
                        'time': event['timestamp']
                    })
                elif event['event_type'] == 'player_join':
                    player_analysis_data[player_name]['join_time'] = event['timestamp']
                elif event['event_type'] == 'player_leave':
                    player_analysis_data[player_name]['leave_time'] = event['timestamp']
        
        # Calculate additional metrics
        for player_name, data in player_analysis_data.items():
            # Calculate total distance traveled
            positions = data['positions']
            if len(positions) > 1:
                total_distance = 0
                for i in range(1, len(positions)):
                    prev = positions[i-1]
                    curr = positions[i]
                    distance = ((curr['x'] - prev['x'])**2 + 
                              (curr['y'] - prev['y'])**2 + 
                              (curr['z'] - prev['z'])**2) ** 0.5
                    total_distance += distance
                data['total_distance'] = round(total_distance, 2)
        
        # Analyze each player
        analyses = {}
        
        for player_name, player_data in player_analysis_data.items():
            # Prepare summary for ChatGPT
            summary = f"""
Player Activity Summary:
- Total positions recorded: {len(player_data['positions'])}
- Total distance traveled: {player_data['total_distance']} blocks
- Blocks placed: {len(player_data['blocks_placed'])}
- Blocks broken: {len(player_data['blocks_broken'])}
- Session duration: {player_data['join_time']} to {player_data['leave_time'] or 'still active'}

Block Placement Details:
"""
            for block in player_data['blocks_placed'][:10]:  # Show first 10
                summary += f"- {block['type']} at ({block['position']['x']}, {block['position']['y']}, {block['position']['z']})\n"
            
            if len(player_data['blocks_placed']) > 10:
                summary += f"... and {len(player_data['blocks_placed']) - 10} more blocks\n"
            
            summary += "\nBlock Breaking Details:\n"
            for block in player_data['blocks_broken'][:10]:  # Show first 10
                summary += f"- {block['type']} at ({block['position']['x']}, {block['position']['y']}, {block['position']['z']})\n"
            
            if len(player_data['blocks_broken']) > 10:
                summary += f"... and {len(player_data['blocks_broken']) - 10} more blocks\n"
            
            # Prepare prompt for ChatGPT
            prompt = f"""
Please analyze the following Minecraft player's gameplay data against the provided rubric.

RUBRIC:
{rubric_content}

PLAYER: {player_name}

{summary}

Please provide a detailed assessment of this player's performance based on the rubric criteria. 
Include specific examples from their gameplay data and suggestions for improvement.
Format the response in a clear, structured way with sections for different rubric criteria.
"""
            
            # Call Azure OpenAI API
            try:
                if client is not None:
                    # New API style
                    response = client.chat.completions.create(
                        model=os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME', 'gpt-4.1'),
                        messages=[
                            {"role": "system", "content": "You are a Minecraft gameplay assessment expert. Analyze player data against the provided rubric and give constructive feedback. Be specific and reference actual gameplay data."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=1000,
                        temperature=0.7
                    )
                    
                    analysis = response.choices[0].message.content
                else:
                    # Old API style fallback
                    import openai
                    response = openai.ChatCompletion.create(
                        engine=os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME', 'gpt-4.1'),
                        messages=[
                            {"role": "system", "content": "You are a Minecraft gameplay assessment expert. Analyze player data against the provided rubric and give constructive feedback. Be specific and reference actual gameplay data."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=1000,
                        temperature=0.7
                    )
                    
                    analysis = response.choices[0].message.content
                
                analyses[player_name] = analysis
                
            except Exception as e:
                analyses[player_name] = f"Error analyzing player: {str(e)}"
        
        # Store the results globally before returning
        latest_assessment_results = {"analyses": analyses}
        return {"analyses": analyses}
        
    except Exception as e:
        logger.error(f"Error in analyze_player_data: {str(e)}")
        return {"error": str(e)}

class SimpleHTTPHandler(BaseHTTPRequestHandler):
    """Simple HTTP handler to serve the static files"""
    
    def do_GET(self):
        # Map paths to files
        static_dir = os.path.join(os.path.dirname(__file__), 'static')
        
        if self.path == '/':
            self.path = '/index.html'
        
        # Special endpoint for server info
        if self.path == '/api/server-info':
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            external_ip = get_external_ip()
            info = {
                'external_ip': external_ip,
                'minecraft_port': 19131,
                'connection_string': f'/connect {external_ip}:19131'
            }
            self.wfile.write(json.dumps(info).encode())
            return
        
        # Special endpoint for session export
        if self.path.startswith('/api/export-session/'):
            session_name = self.path.split('/')[-1]
            session_file_path = os.path.join(DATA_DIR, f"{session_name}.json")
            
            try:
                if os.path.exists(session_file_path):
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json')
                    self.send_header('Content-Disposition', f'attachment; filename="{session_name}.json"')
                    self.end_headers()
                    
                    with open(session_file_path, 'rb') as f:
                        self.wfile.write(f.read())
                    
                    logger.info(f"📁 Exported session file: {session_name}.json")
                    return
                else:
                    self.send_response(404)
                    self.send_header('Content-type', 'application/json')
                    self.end_headers()
                    self.wfile.write(json.dumps({'error': 'Session file not found'}).encode())
                    return
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': str(e)}).encode())
                logger.error(f"Error exporting session: {e}")
                return

        # Special endpoint for rubric
        if self.path == '/api/rubric':
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            
            rubric_path = os.path.join(os.path.dirname(__file__), 'rubric.md')
            try:
                with open(rubric_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                response = {'content': content}
            except FileNotFoundError:
                response = {'content': '# Assessment Rubric\n\nNo rubric file found. Create your rubric here.'}
            except Exception as e:
                response = {'content': f'Error reading rubric: {str(e)}'}
            
            self.wfile.write(json.dumps(response).encode())
            return
        
        # Special endpoint for downloading assessment as Word document
        if self.path == '/api/download-assessment':
            try:
                if not latest_assessment_results.get('analyses'):
                    self.send_response(404)
                    self.send_header('Content-type', 'application/json')
                    self.end_headers()
                    self.wfile.write(json.dumps({'error': 'No assessment results available'}).encode())
                    return
                
                # Create Word document
                doc = create_assessment_document(latest_assessment_results['analyses'])
                
                # Save to bytes buffer
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                doc_bytes = doc_buffer.read()
                
                # Generate filename
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'minecraft_assessment_{timestamp}.docx'
                
                # Send response
                self.send_response(200)
                self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
                self.send_header('Content-Length', str(len(doc_bytes)))
                self.end_headers()
                self.wfile.write(doc_bytes)
                
                logger.info(f"📄 Generated assessment document: {filename}")
                
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': str(e)}).encode())
                logger.error(f"Error generating assessment document: {e}")
            return
        
        # Remove query string if present
        clean_path = self.path.split('?')[0]
        
        # Security: prevent directory traversal
        if '..' in clean_path:
            self.send_response(403)
            self.end_headers()
            return
        
        # Determine content type
        content_types = {
            '.html': 'text/html',
            '.css': 'text/css',
            '.js': 'application/javascript',
            '.json': 'application/json'
        }
        
        file_ext = os.path.splitext(clean_path)[1]
        content_type = content_types.get(file_ext, 'text/plain')
        
        # Try to serve the file
        file_path = os.path.join(static_dir, clean_path.lstrip('/'))
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            self.send_response(200)
            self.send_header('Content-type', content_type)
            self.end_headers()
            self.wfile.write(content)
        except FileNotFoundError:
            self.send_response(404)
            self.end_headers()
            self.wfile.write(b'File not found')
    
    def do_POST(self):
        """Handle POST requests"""
        if self.path == '/api/rubric':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            
            try:
                data = json.loads(post_data.decode('utf-8'))
                content = data.get('content', '')
                
                rubric_path = os.path.join(os.path.dirname(__file__), 'rubric.md')
                
                # Create backup of existing rubric
                if os.path.exists(rubric_path):
                    backup_path = rubric_path + '.backup'
                    with open(rubric_path, 'r', encoding='utf-8') as f:
                        backup_content = f.read()
                    with open(backup_path, 'w', encoding='utf-8') as f:
                        f.write(backup_content)
                
                # Save new rubric
                with open(rubric_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'success': True}).encode())
                
                logger.info("📝 Rubric updated successfully")
                
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': str(e)}).encode())
                logger.error(f"Error saving rubric: {e}")
        else:
            self.send_response(404)
            self.end_headers()
    
    def log_message(self, format, *args):
        # Suppress HTTP logs
        pass

async def broadcast_to_web(message):
    """Broadcast updates to all connected web clients"""
    disconnected = set()
    for client in web_clients:
        try:
            await client.send(json.dumps(message))
        except:
            disconnected.add(client)
    # Remove disconnected clients
    web_clients.difference_update(disconnected)

async def send_message_to_minecraft(websocket, message):
    """Send a chat message to all Minecraft players"""
    # Format the message with color and styling
    formatted_message = f"§e§l[Server]§r §f{message}"
    
    command = {
        "header": {
            "version": 1,
            "requestId": str(uuid4()),
            "messageType": "commandRequest",
            "messagePurpose": "commandRequest"
        },
        "body": {
            "origin": {"type": "player"},
            "commandLine": f'tellraw @a {{"rawtext":[{{"text":"{formatted_message}"}}]}}',
            "version": 1
        }
    }
    
    await websocket.send(json.dumps(command))
    logger.info(f"📢 Sent message to players: {message}")

async def send_command_to_minecraft(websocket, command, target_players=None, is_player_specific=False):
    """Send a command to Minecraft"""
    # Prepare the command
    if is_player_specific and target_players:
        # For player-specific commands, we need to run the command for each player
        for player_name in target_players:
            player_command = command.replace('@t', f'@a[name={player_name}]')
            cmd = {
                "header": {
                    "version": 1,
                    "requestId": str(uuid4()),
                    "messageType": "commandRequest",
                    "messagePurpose": "commandRequest"
                },
                "body": {
                    "origin": {"type": "player"},
                    "commandLine": player_command,
                    "version": 1
                }
            }
            await websocket.send(json.dumps(cmd))
            logger.info(f"🎮 Sent command for {player_name}: {player_command}")
    else:
        # For global commands, send once
        cmd = {
            "header": {
                "version": 1,
                "requestId": str(uuid4()),
                "messageType": "commandRequest",
                "messagePurpose": "commandRequest"
            },
            "body": {
                "origin": {"type": "player"},
                "commandLine": command,
                "version": 1
            }
        }
        await websocket.send(json.dumps(cmd))
        logger.info(f"🎮 Sent command: {command}")

# Store active Minecraft connections
minecraft_connections = set()

async def handle_web_client(websocket):
    """Handle web browser WebSocket connections for live updates"""
    logger.info("Web client connected for live updates")
    web_clients.add(websocket)
    
    try:
        # Send session info
        if session_id:
            await websocket.send(json.dumps({
                'type': 'session_info',
                'sessionId': session_id,
                'startTime': session_start_time.isoformat(),
                'fileName': os.path.basename(session_file)
            }))
        
        # Send current player positions
        for player_id, pos in player_positions.items():
            await websocket.send(json.dumps({
                'type': 'position',
                'playerId': player_id,
                'playerName': pos.get('name', player_id),
                'x': pos['x'],
                'y': pos['y'],
                'z': pos['z']
            }))
        
        # Keep connection alive and handle messages
        async for message in websocket:
            try:
                msg = json.loads(message)
                
                if msg.get('type') == 'analyze_request':
                    # Perform analysis
                    logger.info("Received AI analysis request")
                    analysis_result = await analyze_player_data()
                    
                    # Send result back
                    await websocket.send(json.dumps({
                        'type': 'analysis_result',
                        **analysis_result
                    }))
                
                elif msg.get('type') == 'clear_session':
                    # Handle clear session request
                    logger.info("Received clear session request")
                    
                    # End current session if active
                    if session_id and active_players:
                        end_session()
                    
                    # Clear global data
                    global session_events, event_buffer, last_save_time
                    session_events = []
                    event_buffer = []
                    last_save_time = time.time()
                    
                    # Start new session if there are active players
                    if active_players:
                        start_session()
                        
                        # Send new session info to all web clients
                        await broadcast_to_web({
                            'type': 'session_cleared',
                            'sessionId': session_id,
                            'startTime': session_start_time.isoformat() if session_start_time else None,
                            'fileName': os.path.basename(session_file) if session_file else None
                        })
                    else:
                        # No active players, just reset
                        await broadcast_to_web({
                            'type': 'session_cleared',
                            'sessionId': None,
                            'startTime': None,
                            'fileName': None
                        })
                    
                    logger.info("✨ Session cleared and ready for new recording")
                
                elif msg.get('type') == 'send_message':
                    # Send message to all Minecraft clients
                    message_text = msg.get('message', '')
                    if message_text:
                        # Record the message event
                        record_event("server_message", {
                            "message": message_text,
                            "sender": "web_interface"
                        })
                        
                        # Send to all connected Minecraft clients
                        disconnected = set()
                        for mc_client in minecraft_connections:
                            try:
                                await send_message_to_minecraft(mc_client, message_text)
                            except:
                                disconnected.add(mc_client)
                        
                        # Remove disconnected clients
                        minecraft_connections.difference_update(disconnected)
                        
                        logger.info(f"📨 Message sent to {len(minecraft_connections)} Minecraft clients")
                
                elif msg.get('type') == 'game_command':
                    # Handle game command request
                    command = msg.get('command', '')
                    target_mode = msg.get('targetMode', 'all')
                    target_players = msg.get('targetPlayers', [])
                    is_player_specific = msg.get('isPlayerSpecific', False)
                    
                    if command:
                        # Record the command event
                        record_event("game_command", {
                            "command": command,
                            "target_mode": target_mode,
                            "target_players": target_players,
                            "sender": "web_interface"
                        })
                        
                        # Send to all connected Minecraft clients
                        disconnected = set()
                        
                        # Get player names for selected players
                        selected_player_names = []
                        if target_mode == 'selected' and target_players:
                            for player_id in target_players:
                                for pid, pos in player_positions.items():
                                    if pid == player_id:
                                        selected_player_names.append(pos.get('name', player_id))
                        
                        for mc_client in minecraft_connections:
                            try:
                                if target_mode == 'all' or not is_player_specific:
                                    await send_command_to_minecraft(mc_client, command)
                                else:
                                    await send_command_to_minecraft(mc_client, command, selected_player_names, is_player_specific)
                            except:
                                disconnected.add(mc_client)
                        
                        # Remove disconnected clients
                        minecraft_connections.difference_update(disconnected)
                        
                        logger.info(f"🎮 Command sent to {len(minecraft_connections)} Minecraft clients")
                        
            except Exception as e:
                logger.error(f"Error handling web client message: {e}")
            
    except websockets.exceptions.ConnectionClosed:
        pass
    finally:
        web_clients.discard(websocket)
        logger.info("Web client disconnected")

async def send_welcome_message(websocket):
    """Send welcome message to all players when they join"""
    welcome_command = {
        "header": {
            "version": 1,
            "requestId": str(uuid4()),
            "messageType": "commandRequest",
            "messagePurpose": "commandRequest"
        },
        "body": {
            "origin": {"type": "player"},
            "commandLine": 'tellraw @a {"rawtext":[{"text":"§6§l======\\n§r§e§lWelcome to Playtrace AI 1.1\\n§r§eYour game data is being recorded.\\n§eIf you do not want this please exit now.\\n§6§l======"}]}',
            "version": 1
        }
    }
    await websocket.send(json.dumps(welcome_command))
    logger.info(f"📢 Sent welcome message")

async def handle_minecraft_client(websocket):
    """Handle Minecraft Education Edition connections"""
    global active_players
    
    # Add this connection to the set
    minecraft_connections.add(websocket)
    
    try:
        client_addr = f"{websocket.remote_address[0]}:{websocket.remote_address[1]}"
    except:
        client_addr = "unknown"
        
    logger.info(f"🎮 Minecraft connected from {client_addr}")
    
    player_id = None
    player_name = None
    welcome_sent = False
    
    # Start session when first player connects
    if len(active_players) == 0:
        start_session()
        # Broadcast session info to web clients
        await broadcast_to_web({
            'type': 'session_info',
            'sessionId': session_id,
            'startTime': session_start_time.isoformat(),
            'fileName': os.path.basename(session_file)  # Send just the filename
        })
    
    try:
        # Send initial connection confirmation
        await websocket.send(json.dumps({
            "header": {"messagePurpose": "commandResponse"},
            "body": {"statusMessage": f"Connected to Playtrace AI! Session: {os.path.basename(session_file)}"}
        }))
        
        # Subscribe to a comprehensive list of events
        events_to_subscribe = [
            # Core events we already have
            "BlockPlaced", "BlockBroken", "PlayerTravelled", "PlayerMessage",
            # Additional player events
            "ItemUsed", "ItemInteracted", "ItemCrafted", "ItemSmelted",
            "ItemEquipped", "ItemDropped", "ItemPickedUp",
            # Combat and damage events
            "PlayerDied", "MobKilled", "PlayerHurt", "PlayerAttack",
            # World interaction events
            "DoorUsed", "ChestOpened", "ContainerClosed", "ButtonPressed",
            "LeverUsed", "PressurePlateActivated", 
            # Movement events
            "PlayerJump", "PlayerSneak", "PlayerSprint", "PlayerSwim",
            "PlayerClimb", "PlayerGlide", "PlayerTeleport",
            # Achievement/advancement events
            "AwardAchievement", "PlayerTransform",
            # Entity events
            "EntitySpawned", "EntityRemoved", "EntityInteracted",
            # World events
            "WeatherChanged", "TimeChanged", "GameRulesUpdated",
            # Other useful events
            "PlayerEat", "PlayerSleep", "PlayerWake", "CameraUsed",
            "BookEdited", "BossKilled", "RaidCompleted", "TradeCompleted"
        ]
        
        for event_name in events_to_subscribe:
            await websocket.send(json.dumps({
                "header": {
                    "version": 1,
                    "requestId": str(uuid4()),
                    "messageType": "commandRequest",
                    "messagePurpose": "subscribe"
                },
                "body": {
                    "eventName": event_name
                }
            }))
        
        logger.info(f"📋 Subscribed to {len(events_to_subscribe)} event types")
        
        # Position update counter for throttling
        position_update_counter = 0
        
        async for message in websocket:
            try:
                msg = json.loads(message)
                
                # Extract header and body
                header = msg.get('header', {})
                body = msg.get('body', {})
                event_name = header.get('eventName', '')
                
                # Get player info if available
                if 'player' in body and player_id is None:
                    player_data = body['player']
                    player_id = str(player_data.get('id', f"Player_{int(time.time()) % 1000}"))
                    player_name = player_data.get('name', player_id)
                    active_players.add(player_id)
                    logger.info(f"🎯 Identified player: {player_name}")
                    
                    # Send welcome message
                    if not welcome_sent:
                        await send_welcome_message(websocket)
                        welcome_sent = True
                    
                    # Record player join event
                    record_event("player_join", {
                        "player_id": player_id,
                        "player_name": player_name,
                        "address": client_addr
                    })
                
                # Handle all events
                if header.get('messagePurpose') == 'event':
                    # Extract common player data
                    player_data = body.get('player', {})
                    current_player_id = str(player_data.get('id', player_id)) if player_data else player_id
                    current_player_name = player_data.get('name', player_name) if player_data else player_name
                    
                    # Handle PlayerMessage event for chat
                    if event_name == 'PlayerMessage':
                        message_type = body.get('type', '')
                        message_text = body.get('message', '')
                        sender = body.get('sender', 'Unknown')
                        
                        # Record chat event
                        record_event("player_chat", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "message": message_text,
                            "message_type": message_type,
                            "sender": sender
                        })
                        
                        # Broadcast to web clients
                        await broadcast_to_web({
                            'type': 'player_chat',
                            'playerId': current_player_id,
                            'playerName': current_player_name,
                            'message': message_text
                        })
                        
                        logger.info(f"💬 Chat from {current_player_name}: {message_text}")
                    
                    # Handle PlayerTravelled event
                    elif event_name == 'PlayerTravelled':
                        if 'position' in player_data:
                            pos = player_data['position']
                            pos_x = float(pos['x'])
                            pos_y = float(pos['y'])
                            pos_z = float(pos['z'])
                            
                            player_positions[current_player_id] = {
                                'x': pos_x,
                                'y': pos_y,
                                'z': pos_z,
                                'name': current_player_name
                            }
                            
                            # Only record position every 10 updates to avoid spam
                            position_update_counter += 1
                            if (position_update_counter % 10) == 0:
                                record_event("player_position", {
                                    "player_id": current_player_id,
                                    "player_name": current_player_name,
                                    "position": {"x": pos_x, "y": pos_y, "z": pos_z},
                                    "dimension": player_data.get('dimension', 'overworld')
                                })
                        
                        await broadcast_to_web({
                            'type': 'position',
                            'playerId': current_player_id,
                            'playerName': current_player_name,
                            'x': pos_x,
                            'y': pos_y,
                            'z': pos_z
                        })
                    
                    # Handle block events (existing code)
                    elif event_name == 'BlockPlaced':
                        player_pos = body.get('player', {}).get('position', {})
                        block_x = int(player_pos.get('x', 0))
                        block_y = int(player_pos.get('y', 0))
                        block_z = int(player_pos.get('z', 0))
                        
                        block_info = body.get('block', {})
                        block_id = block_info.get('id', 'unknown')
                        block_namespace = block_info.get('namespace', 'minecraft')
                        block_type = f"{block_namespace}:{block_id}"
                        
                        # Record block placed event
                        record_event("block_placed", {
                            "player_id": player_id,
                            "player_name": player_name,
                            "block_type": block_type,
                            "player_position": {"x": player_pos.get('x', 0), "y": player_pos.get('y', 0), "z": player_pos.get('z', 0)},
                            "estimated_block_position": {"x": block_x, "y": block_y + 1, "z": block_z}
                        })
                        
                        await broadcast_to_web({
                            'type': 'block_place',
                            'x': player_pos.get('x', 0),
                            'y': player_pos.get('y', 0),
                            'z': player_pos.get('z', 0),
                            'blockPos': {
                                'x': block_x,
                                'y': block_y + 1,
                                'z': block_z
                            },
                            'blockType': block_type,
                            'playerName': player_name or 'Unknown'
                        })
                        
                        logger.info(f"🟩 Block placed: {block_type} near ({block_x}, {block_y}, {block_z}) by {player_name}")
                    
                    elif event_name == 'BlockBroken':
                        player_pos = body.get('player', {}).get('position', {})
                        block_x = int(player_pos.get('x', 0))
                        block_y = int(player_pos.get('y', 0))
                        block_z = int(player_pos.get('z', 0))
                        
                        block_info = body.get('block', {})
                        block_id = block_info.get('id', 'unknown')
                        block_namespace = block_info.get('namespace', 'minecraft')
                        block_type = f"{block_namespace}:{block_id}"
                        
                        # Record block broken event
                        record_event("block_broken", {
                            "player_id": player_id,
                            "player_name": player_name,
                            "block_type": block_type,
                            "player_position": {"x": player_pos.get('x', 0), "y": player_pos.get('y', 0), "z": player_pos.get('z', 0)},
                            "estimated_block_position": {"x": block_x, "y": block_y, "z": block_z}
                        })
                        
                        await broadcast_to_web({
                            'type': 'block_break',
                            'x': player_pos.get('x', 0),
                            'y': player_pos.get('y', 0),
                            'z': player_pos.get('z', 0),
                            'blockPos': {
                                'x': block_x,
                                'y': block_y,
                                'z': block_z
                            },
                            'blockType': block_type,
                            'playerName': player_name or 'Unknown'
                        })
                        
                        logger.info(f"🟥 Block broken: {block_type} near ({block_x}, {block_y}, {block_z}) by {player_name}")
                    
                    # Handle item events
                    elif event_name == 'ItemUsed':
                        item_data = body.get('item', {})
                        record_event("item_used", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "item": f"{item_data.get('namespace', 'minecraft')}:{item_data.get('id', 'unknown')}",
                            "count": body.get('count', 1),
                            "position": player_data.get('position', {})
                        })
                        logger.info(f"🔧 {current_player_name} used item: {item_data.get('id', 'unknown')}")
                    
                    elif event_name == 'ItemCrafted':
                        item_data = body.get('item', {})
                        record_event("item_crafted", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "item": f"{item_data.get('namespace', 'minecraft')}:{item_data.get('id', 'unknown')}",
                            "count": body.get('count', 1)
                        })
                        logger.info(f"🔨 {current_player_name} crafted: {item_data.get('id', 'unknown')}")
                    
                    elif event_name == 'ItemEquipped':
                        item_data = body.get('item', {})
                        record_event("item_equipped", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "item": f"{item_data.get('namespace', 'minecraft')}:{item_data.get('id', 'unknown')}",
                            "slot": body.get('slot', 'unknown')
                        })
                    
                    elif event_name == 'ItemDropped':
                        item_data = body.get('item', {})
                        record_event("item_dropped", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "item": f"{item_data.get('namespace', 'minecraft')}:{item_data.get('id', 'unknown')}",
                            "count": body.get('count', 1),
                            "position": player_data.get('position', {})
                        })
                    
                    elif event_name == 'ItemPickedUp':
                        item_data = body.get('item', {})
                        record_event("item_picked_up", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "item": f"{item_data.get('namespace', 'minecraft')}:{item_data.get('id', 'unknown')}",
                            "count": body.get('count', 1),
                            "position": player_data.get('position', {})
                        })
                    
                    # Handle combat events
                    elif event_name == 'PlayerDied':
                        death_cause = body.get('cause', 'unknown')
                        killer = body.get('killer', {})
                        record_event("player_died", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "cause": death_cause,
                            "killer": killer.get('name', 'unknown') if killer else 'environment',
                            "position": player_data.get('position', {})
                        })
                        logger.info(f"💀 {current_player_name} died: {death_cause}")
                        
                        # Broadcast death event
                        await broadcast_to_web({
                            'type': 'player_event',
                            'eventType': 'death',
                            'playerId': current_player_id,
                            'playerName': current_player_name,
                            'details': f"died from {death_cause}"
                        })
                    
                    elif event_name == 'MobKilled':
                        mob_data = body.get('mob', {})
                        record_event("mob_killed", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "mob_type": mob_data.get('id', 'unknown'),
                            "mob_name": mob_data.get('name', ''),
                            "weapon": body.get('weapon', 'unknown'),
                            "position": player_data.get('position', {})
                        })
                        logger.info(f"⚔️ {current_player_name} killed {mob_data.get('id', 'mob')}")
                    
                    elif event_name == 'PlayerHurt':
                        damage = body.get('damage', 0)
                        cause = body.get('cause', 'unknown')
                        record_event("player_hurt", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "damage": damage,
                            "cause": cause,
                            "health": player_data.get('health', 0),
                            "position": player_data.get('position', {})
                        })
                    
                    # Handle interaction events
                    elif event_name == 'DoorUsed':
                        door_data = body.get('door', {})
                        record_event("door_used", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "door_type": door_data.get('id', 'unknown'),
                            "action": body.get('action', 'interact'),
                            "position": body.get('block_position', player_data.get('position', {}))
                        })
                    
                    elif event_name == 'ChestOpened':
                        chest_data = body.get('chest', {})
                        record_event("chest_opened", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "chest_type": chest_data.get('id', 'chest'),
                            "position": body.get('block_position', player_data.get('position', {}))
                        })
                        logger.info(f"📦 {current_player_name} opened chest")
                    
                    # Handle movement events
                    elif event_name in ['PlayerJump', 'PlayerSneak', 'PlayerSprint', 'PlayerSwim', 'PlayerClimb', 'PlayerGlide']:
                        record_event(f"player_{event_name.lower().replace('player', '')}", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "position": player_data.get('position', {}),
                            "duration": body.get('duration', 0)
                        })
                    
                    # Handle achievements
                    elif event_name == 'AwardAchievement':
                        achievement = body.get('achievement', 'unknown')
                        record_event("achievement_earned", {
                            "player_id": current_player_id,
                            "player_name": current_player_name,
                            "achievement": achievement,
                            "description": body.get('description', '')
                        })
                        logger.info(f"🏆 {current_player_name} earned achievement: {achievement}")
                        
                        # Broadcast achievement
                        await broadcast_to_web({
                            'type': 'player_event',
                            'eventType': 'achievement',
                            'playerId': current_player_id,
                            'playerName': current_player_name,
                            'details': f"earned {achievement}"
                        })
                    
                    # Handle other events generically
                    else:
                        # Record any other event with full body data
                        if event_name and event_name not in ['PlayerTravelled']:  # Skip high-frequency events
                            record_event(event_name.lower(), {
                                "player_id": current_player_id,
                                "player_name": current_player_name,
                                "event_data": body
                            })
                
                # Check if we need to save
                if len(event_buffer) > 0 and (time.time() - last_save_time > 5):
                    await broadcast_to_web({'type': 'save_notification'})
                
            except Exception as e:
                logger.error(f"Error processing message: {e}")
                logger.error(f"Message content: {message}")
    
    except websockets.exceptions.ConnectionClosed:
        logger.info(f"🎮 Minecraft disconnected: {player_name or 'Unknown'}")
    finally:
        # Clean up player
        if player_id:
            # Record player leave event
            record_event("player_leave", {
                "player_id": player_id,
                "player_name": player_name
            })
            
            if player_id in player_positions:
                del player_positions[player_id]
            
            if player_id in active_players:
                active_players.remove(player_id)
            
            await broadcast_to_web({
                'type': 'disconnect',
                'playerId': player_id
            })
            
            # End session when last player leaves
            if len(active_players) == 0:
                logger.info("📤 Last player left, ending session...")
                end_session()

def run_http_server():
    """Run HTTP server in a separate thread"""
    server = HTTPServer(('0.0.0.0', 8080), SimpleHTTPHandler)  # Bind to all interfaces
    logger.info("🌐 Web interface running at http://localhost:8080")
    server.serve_forever()

async def main():
    """Main function to run both servers"""
    logger.info("=" * 60)
    logger.info("🎮 Minecraft 3D Live Tracker with AI Assessment")
    logger.info("=" * 60)
    
    # Ensure data directory exists
    ensure_data_directory()
    
    # Get external IP
    external_ip = get_external_ip()
    
    # Check for Azure OpenAI API key
    if not os.getenv('AZURE_OPENAI_API_KEY'):
        logger.warning("⚠️  Azure OpenAI API key not found! AI assessment will not work.")
        logger.warning("   Set AZURE_OPENAI_API_KEY in your .env file")
    
    # Start HTTP server in background thread
    http_thread = threading.Thread(target=run_http_server, daemon=True)
    http_thread.start()
    
    # Start WebSocket servers
    minecraft_server = await websockets.serve(
        handle_minecraft_client, 
        '0.0.0.0',  # Bind to all interfaces
        19131  # Minecraft port
    )
    
    web_server = await websockets.serve(
        handle_web_client,
        '0.0.0.0',  # Bind to all interfaces
        8081  # Web updates port
    )
    
    logger.info("✅ Servers started successfully!")
    logger.info("")
    logger.info(f"📡 Minecraft: Connect with /connect {external_ip}:19131")
    logger.info(f"🌐 3D Viewer: Open http://{external_ip}:8080 in your browser")
    logger.info(f"   Local: http://localhost:8080")
    logger.info("🤖 AI: Click 'Analyze Players with AI' button")
    logger.info("💾 JSON files saved to: {DATA_DIR}/ directory")
    logger.info("=" * 60)
    
    # Run forever
    await asyncio.Future()

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("\n👋 Server stopped")
        # Save any active session
        if active_players:
            logger.info("💾 Saving active session before shutdown...")
            end_session()